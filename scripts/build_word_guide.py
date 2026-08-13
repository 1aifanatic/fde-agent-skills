from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "FDE_AGENT_SKILLS_HANDBOOK.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "19324D"
INK = "222222"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"
WHITE = "FFFFFF"
GOLD = "A66B00"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

SOURCES = [
    ("Part I — Install and verify", ROOT / "docs" / "INSTALLATION.md"),
    ("Part II — Operating model and usage", ROOT / "docs" / "USER_GUIDE.md"),
    ("Part III — Before-and-after scenarios", ROOT / "docs" / "SCENARIOS.md"),
    ("Part IV — Worked example overview", ROOT / "examples" / "northstar-ap-transformation" / "README.md"),
    ("Part V — Skill execution evidence", ROOT / "examples" / "northstar-ap-transformation" / "run-log.md"),
    ("Part VI — Demonstrated impact", ROOT / "examples" / "northstar-ap-transformation" / "impact.md"),
    ("Part VII — Validation and reproducibility", ROOT / "docs" / "VALIDATION.md"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_dxa / 1440)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    grid_cols = grid.findall(qn("w:gridCol"))
    while len(grid_cols) < len(widths_dxa):
        grid.append(OxmlElement("w:gridCol"))
        grid_cols = grid.findall(qn("w:gridCol"))
    for grid_col, width in zip(grid_cols, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def distribute_widths(rows: list[list[str]]) -> list[int]:
    column_count = max(len(row) for row in rows)
    scores = []
    for index in range(column_count):
        values = [row[index] if index < len(row) else "" for row in rows]
        max_len = max(len(re.sub(r"\[[^\]]+\]\([^)]+\)", "link", value)) for value in values)
        scores.append(max(9, min(max_len, 46)))

    total_score = sum(scores)
    raw = [max(1000, round(CONTENT_WIDTH_DXA * score / total_score)) for score in scores]
    scale = CONTENT_WIDTH_DXA / sum(raw)
    widths = [max(900, round(value * scale)) for value in raw]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference

    if widths[-1] < 900:
        deficit = 900 - widths[-1]
        donor = max(range(len(widths) - 1), key=lambda i: widths[i])
        widths[donor] -= deficit
        widths[-1] = 900
    return widths


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, size: float | None = None, color: str | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            if size is not None or color is not None:
                set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            if size is not None or color is not None:
                set_run_font(run, size=size, color=color)
            run.bold = True
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if size is not None or color is not None:
            set_run_font(run, size=size, color=color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
        "Heading 4": (11, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("FDE AGENT SKILLS HANDBOOK")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run("\tPRIVATE REPOSITORY GUIDE")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    label = paragraph.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    of_run = paragraph.add_run(" of ")
    set_run_font(of_run, size=8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("FORWARD DEPLOYED ENGINEERING")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("FDE Agent Skills")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Installation, Operating Model, Scenarios, and a Fully Worked Engagement")
    set_run_font(run, size=15, color=DARK_BLUE)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Inches(0.65)
    statement.paragraph_format.right_indent = Inches(0.65)
    statement.paragraph_format.space_after = Pt(32)
    run = statement.add_run(
        "A durable, evidence-backed way for an FDE and an AI agent to understand a customer, "
        "redesign work, plan delivery, and control change—without a custom web application."
    )
    set_run_font(run, size=11.5, color=INK)

    for label, value in (
        ("Repository", "github.com/1aifanatic/fde-agent-skills (private)"),
        ("Suite", "6 interoperable skills and a 21-file engagement workspace"),
        ("Edition", f"Generated {date.today().isoformat()}"),
        ("Demonstration", "Synthetic Northstar AP transformation; no production claims"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        lead = p.add_run(f"{label}: ")
        set_run_font(lead, size=9.5, color=MUTED, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.left_indent = Inches(0.7)
    note.paragraph_format.right_indent = Inches(0.7)
    run = note.add_run(
        "Confidentiality note: live customer workspaces may contain sensitive evidence. "
        "Follow company policy and never commit secrets or unnecessary personal data."
    )
    set_run_font(run, size=8.8, color=MUTED, italic=True)

    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("How to use this handbook", level=1)
    p = doc.add_paragraph()
    add_inline(
        p,
        "Read Part I to install the skills. Use Part II while running engagements. "
        "Use Part III to explain the value to stakeholders. Parts IV–VI show the complete synthetic execution and its impact. "
        "Part VII documents the reproducibility and validation evidence."
    )

    entries = [
        ("Part I", "Install and verify the private repository"),
        ("Part II", "Understand the operating model and use every skill"),
        ("Part III", "Compare before-and-after scenarios"),
        ("Part IV", "Orient to the synthetic Northstar AP engagement"),
        ("Part V", "Review the skill-by-skill execution record"),
        ("Part VI", "Understand demonstrated versus hypothesized impact"),
        ("Part VII", "Reproduce installation, skill, example, and document checks"),
    ]
    for label, description in entries:
        p = doc.add_paragraph(style="List Bullet")
        lead = p.add_run(f"{label}: ")
        set_run_font(lead, size=11, color=NAVY, bold=True)
        run = p.add_run(description)
        set_run_font(run, size=11, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    lead = p.add_run("Fastest start: ")
    set_run_font(lead, bold=True, color=NAVY)
    add_inline(
        p,
        "Authenticate with GitHub CLI, run the one-command global Codex install, reload Codex, "
        "and invoke $fde-run-engagement."
    )
    doc.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    widths = distribute_widths(normalized)
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                tr_pr = table.rows[0]._tr.get_or_add_trPr()
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                tr_pr.append(repeat)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, size=8.6)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    paragraph._p.get_or_add_pPr().append(shading)
    for index, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=NAVY)
        if index < len(code_lines) - 1:
            run.add_break()


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F9FC")
    paragraph._p.get_or_add_pPr().append(shading)
    add_inline(paragraph, text, size=10.5, color=DARK_BLUE)


def strip_heading_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text).strip()


def render_markdown(doc: Document, path: Path, part_title: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)

    doc.add_heading(part_title, level=1)
    source = path.relative_to(ROOT).as_posix()
    source_p = doc.add_paragraph()
    source_p.paragraph_format.space_after = Pt(10)
    lead = source_p.add_run("Maintained source: ")
    set_run_font(lead, size=8.5, color=MUTED, bold=True)
    run = source_p.add_run(source)
    set_run_font(run, size=8.5, color=MUTED)

    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    skipped_first_h1 = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("~~~"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            title = strip_heading_links(heading.group(2))
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
            else:
                doc.add_heading(title, level=min(level, 4))
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            add_quote(doc, stripped[1:].strip())
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()
    if code_lines:
        add_code_block(doc, code_lines)
    doc.add_page_break()


def add_document_notes(doc: Document) -> None:
    doc.add_heading("Document status and interpretation", level=1)
    points = [
        "The Northstar engagement is synthetic. Names, systems, measures, emails, approvals, and outcomes are fictional.",
        "A design target is not observed ROI. The example deliberately keeps implementation, UAT, release, and production measurement open.",
        "The skills strengthen engagement quality; they do not grant production authority or replace domain, security, test, or release owners.",
        "For the authoritative current skill instructions, read the SKILL.md files in the repository.",
    ]
    for point in points:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_inline(paragraph, point)

    doc.add_heading("Upstream references", level=2)
    for label, url in (
        ("skills CLI documentation", "https://www.skills.sh/docs/cli"),
        ("skills CLI source", "https://github.com/vercel-labs/skills"),
        ("Codex use cases", "https://developers.openai.com/codex/use-cases"),
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        add_hyperlink(paragraph, label, url)

    doc.add_heading("Repository", level=2)
    paragraph = doc.add_paragraph()
    add_hyperlink(paragraph, "github.com/1aifanatic/fde-agent-skills", "https://github.com/1aifanatic/fde-agent-skills")


def build() -> Path:
    doc = Document()
    configure_page(doc)
    configure_styles(doc)

    properties = doc.core_properties
    properties.title = "FDE Agent Skills Handbook"
    properties.subject = "Installation, operating model, scenarios, and worked example"
    properties.author = "FDE Agent Skills"
    properties.keywords = "forward deployed engineering, FDE, AI agents, skills, process redesign, delivery governance"
    properties.comments = "Generated from repository Markdown sources."

    add_cover(doc)
    add_contents(doc)

    for part_title, source in SOURCES:
        render_markdown(doc, source, part_title)

    add_document_notes(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
